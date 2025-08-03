import asyncio
import os
import sys
import shutil

from typing import Dict, Any
from fastmcp import FastMCP
import aiofiles
import os
import asyncio
from typing import Dict, Any
from docx import Document

# Print startup message
print("[fastfs-mcp] Server starting...", file=sys.stderr, flush=True)

# Initialize the MCP server
mcp = FastMCP(name="fastfs-mcp")

@mcp.tool(description="Move or rename files or directories.")
def mv(source: str, destination: str) -> str:
    """Move or rename files or directories."""
    try:
        print(f"[DEBUG] mv called with source: {source}, destination: {destination}", file=sys.stderr, flush=True)
        if not os.path.exists(source):
            return f"Error: Source '{source}' does not exist"
        
        shutil.move(source, destination)
        return f"Successfully moved '{source}' to '{destination}'"
    except Exception as e:
        print(f"[ERROR] mv failed: {str(e)}", file=sys.stderr, flush=True)
        return f"Error: {str(e)}"

@mcp.tool(description="Remove files or directories.")
def rm(path: str, recursive: bool = False, force: bool = False) -> str:
    """Remove files or directories."""
    try:
        print(f"[DEBUG] rm called with path: {path}", file=sys.stderr, flush=True)
        if not os.path.exists(path):
            if force:
                return f"Warning: Path '{path}' does not exist, nothing removed"
            else:
                return f"Error: Path '{path}' does not exist"
        
        if os.path.isdir(path):
            if not recursive:
                return f"Error: '{path}' is a directory, use recursive=True to remove directories"
            shutil.rmtree(path)
            return f"Successfully removed directory '{path}'"
        else:
            os.remove(path)
            return f"Successfully removed file '{path}'"
    except Exception as e:
        print(f"[ERROR] rm failed: {str(e)}", file=sys.stderr, flush=True)
        return f"Error: {str(e)}"

@mcp.tool(description="Create a new directory.")
def mkdir(path: str, parents: bool = False) -> str:
    """Create a new directory."""
    try:
        print(f"[DEBUG] mkdir called with path: {path}", file=sys.stderr, flush=True)
        if os.path.exists(path):
            return f"Error: Path '{path}' already exists"
        
        if parents:
            os.makedirs(path)
        else:
            os.mkdir(path)
        return f"Successfully created directory '{path}'"
    except Exception as e:
        print(f"[ERROR] mkdir failed: {str(e)}", file=sys.stderr, flush=True)
        return f"Error: {str(e)}"

# --- ASYNC HELPER ---
async def _run_in_executor(func, *args) -> Any:
    """Запускает блокирующую функцию в отдельном потоке, чтобы не блокировать event loop."""
    return await asyncio.get_running_loop().run_in_executor(None, func, *args)

# ===== ASYNC Word (DOCX) FUNCTIONS =====
@mcp.tool(description="Creates a new Word document (.docx).")
async def create_docx(file_path: str, content: str = "") -> Dict[str, Any]:
    """Асинхронно создает новый документ Word (.docx)."""
    def worker():
        try:
            document = Document()
            if content:
                document.add_paragraph(content)
            document.save(file_path)
            return {"status": "success", "message": f"DOCX файл '{file_path}' успешно создан."}
        except Exception as e:
            return {"status": "error", "message": f"Ошибка при создании DOCX файла: {e}"}
    
    return await _run_in_executor(worker)

@mcp.tool(description="Adds a new paragraph to a DOCX document.")
async def add_text_to_docx(file_path: str, text_to_add: str) -> Dict[str, Any]:
    """Асинхронно добавляет новый абзац в DOCX-документ."""
    def worker():
        try:
            if not os.path.exists(file_path):
                return {"status": "error", "message": f"DOCX файл '{file_path}' не найден."}
            
            document = Document(file_path)
            document.add_paragraph(text_to_add)
            document.save(file_path)
            return {"status": "success", "message": f"Текст успешно добавлен в '{file_path}'."}
        except Exception as e:
            return {"status": "error", "message": f"Ошибка при добавлении текста в DOCX: {e}"}

    return await _run_in_executor(worker)

@mcp.tool(description="Replaces text in a DOCX document.")
async def replace_text_in_docx(file_path: str, old_text: str, new_text: str) -> Dict[str, Any]:
    """Асинхронно заменяет текст в DOCX-документе."""
    def worker():
        try:
            if not os.path.exists(file_path):
                return {"status": "error", "message": f"DOCX файл '{file_path}' не найден."}
            
            document = Document(file_path)
            replaced_count = 0
            for p in document.paragraphs:
                if old_text in p.text:
                    inline = p.runs
                    # Замена текста с сохранением форматирования
                    for i in range(len(inline)):
                        if old_text in inline[i].text:
                            text = inline[i].text.replace(old_text, new_text)
                            inline[i].text = text
                            replaced_count += 1
            
            document.save(file_path)
            return {"status": "success", "message": f"Успешно заменено {replaced_count} вхождений.", "replaced_count": replaced_count}
        except Exception as e:
            return {"status": "error", "message": f"Ошибка при замене текста в DOCX: {e}"}

    return await _run_in_executor(worker)
